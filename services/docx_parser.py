"""
Word document (.docx) parser.

Two distinct use cases:

  1. Content documents — SOPs, manuals, user guides, policies uploaded as
     training content, exactly like PDFs. extract_pages() turns the .docx
     into PageContent objects (the same dataclass PDFParser produces), so
     the file flows through the existing DocumentChunker /
     SemanticChunkingService / Qdrant pipeline unchanged.

  2. Narrative Script Templates — a Word doc whose HEADING STRUCTURE (not
     its body content) defines the exact section layout that a generated
     training script must follow. extract_template_structure() walks the
     template's paragraph styles ("Heading 1", "Heading 2", ...) and
     returns an ordered, indented outline string. The template's own
     narration/body text is intentionally never captured — only the
     headings, since the template defines FORMAT, not factual content.

Docx has no reliable notion of "pages" in its XML (pagination is a
rendering-time concern), so content documents are extracted as a single
PageContent per file. The same structural chunker that splits a PDF page
into headings/sections/paragraphs runs on that text and still detects
section breaks from the heading lines.
"""

import logging
import os
from typing import List

from docx import Document

from services.pdf_parser import PageContent

logger = logging.getLogger(__name__)

_HEADING_STYLE_PREFIX = "Heading"


class DocxParser:
    """Extracts content and structure from .docx files."""

    def __init__(self, docx_folder: str):
        self.docx_folder = docx_folder

    def _list_docx_files(self) -> List[str]:
        return [
            os.path.join(self.docx_folder, f)
            for f in os.listdir(self.docx_folder)
            if f.lower().endswith(".docx")
        ]

    # ------------------------------------------------------------------
    # Content documents (used like PDFs — ingested, chunked, embedded)
    # ------------------------------------------------------------------
    def extract_pages(self, file_path: str) -> List[PageContent]:
        """Read a .docx as training content and return it as PageContent(s)."""
        filename = os.path.basename(file_path)
        doc = Document(file_path)

        lines: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Tables often carry real content in SOPs/manuals — flatten them too
        # so that information isn't silently dropped.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

        full_text = "\n".join(lines)
        if not full_text.strip():
            logger.warning("No extractable text found in '%s'.", filename)
            return []

        core = doc.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "toc_section": "",
        }

        return [
            PageContent(
                filename=filename,
                page_number=1,
                page_label="Document",
                text=full_text,
                metadata=metadata,
            )
        ]

    def extract_all(self) -> List[PageContent]:
        all_pages: List[PageContent] = []
        docx_files = self._list_docx_files()
        logger.info("Found %d .docx file(s) in '%s'", len(docx_files), self.docx_folder)

        for path in docx_files:
            try:
                extracted = self.extract_pages(path)
                logger.info("Extracted %d page(s) from '%s'", len(extracted), os.path.basename(path))
                all_pages.extend(extracted)
            except Exception as exc:
                logger.error("Failed to process '%s': %s", path, exc)

        return all_pages

    # ------------------------------------------------------------------
    # Narrative script templates (structure only — never content)
    # ------------------------------------------------------------------
    @staticmethod
    def extract_template_structure(file_path: str) -> str:
        """Return an ordered, indented outline of the template's headings.

        Example output:
            - Training Title
              - Objectives
                - Learning Outcomes
              - Agenda

        Only heading paragraph styles ("Heading 1", "Heading 2", ... and
        Word's built-in "Title" style) are used to build the outline. Plain
        body paragraphs are ignored, since the template defines section
        structure only — its narration text must never leak into the
        generated script.
        """
        doc = Document(file_path)
        outline_lines: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or not para.style or not para.style.name:
                continue

            style = para.style.name
            if style == "Title":
                level = 0
            elif style.startswith(_HEADING_STYLE_PREFIX):
                digits = "".join(ch for ch in style if ch.isdigit())
                level = (int(digits) - 1) if digits else 0
            else:
                continue

            indent = "  " * max(level, 0)
            outline_lines.append(f"{indent}- {text}")

        if not outline_lines:
            logger.warning(
                "No heading styles found in template '%s' — falling back to "
                "non-empty paragraph text as a flat outline.",
                os.path.basename(file_path),
            )
            outline_lines = [
                f"- {para.text.strip()}"
                for para in doc.paragraphs
                if para.text.strip()
            ]

        return "\n".join(outline_lines)